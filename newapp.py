import io
import re
import os
import json
import requests
from collections import Counter

import streamlit as st

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:
    TfidfVectorizer = None
    cosine_similarity = None


st.set_page_config(
    page_title="AI Resume Rewriter",
    page_icon="resume",
    layout="wide",
)


STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "has",
    "have", "in", "is", "it", "of", "on", "or", "that", "the", "this", "to",
    "with", "you", "your", "we", "our", "will", "can", "using", "use", "into",
    "within", "across", "their", "they", "job", "role", "candidate", "work",
    "experience", "skills", "team", "teams", "responsibilities", "required",
}


def clean_text(text: str) -> str:
    text = re.sub(r"\r", "\n", text or "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_text_from_pdf(file) -> str:
    if PdfReader is None:
        return ""
    reader = PdfReader(file)
    pages = [page.extract_text() or "" for page in reader.pages]
    return clean_text("\n".join(pages))


def extract_text_from_docx(file) -> str:
    if Document is None:
        return ""

    file_bytes = file.getvalue()
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())

    return clean_text("\n".join(parts))


def extract_uploaded_text(file) -> str:
    if file is None:
        return ""

    name = file.name.lower()

    if name.endswith(".txt"):
        return clean_text(file.read().decode("utf-8", errors="ignore"))

    if name.endswith(".pdf"):
        return extract_text_from_pdf(file)

    if name.endswith(".docx"):
        return extract_text_from_docx(file)

    return ""


def tokenize(text: str) -> list[str]:
    words = re.findall(r"[a-zA-Z][a-zA-Z+#.\-]{1,}", text.lower())
    return [w.strip(".-") for w in words if w not in STOPWORDS and len(w) > 2]


def top_keywords(text: str, limit: int = 24) -> list[str]:
    tokens = tokenize(text)
    counts = Counter(tokens)

    phrases = re.findall(
        r"\b(?:machine learning|deep learning|natural language processing|data science|"
        r"large language models|neural networks|computer vision|model deployment|"
        r"streamlit|tensorflow|pytorch|scikit-learn|python|sql|nlp|resume parsing|"
        r"keyword extraction|cosine similarity|text preprocessing|feature extraction)\b",
        text.lower(),
    )

    for phrase in phrases:
        counts[phrase] += 4

    return [word for word, _ in counts.most_common(limit)]


def calculate_match_score(resume_text: str, jd_text: str) -> float:
    if not resume_text or not jd_text:
        return 0.0

    if TfidfVectorizer and cosine_similarity:
        vectorizer = TfidfVectorizer(
            stop_words="english",
            ngram_range=(1, 2),
            max_features=800,
        )
        matrix = vectorizer.fit_transform([resume_text, jd_text])
        return round(float(cosine_similarity(matrix[0], matrix[1])[0][0]) * 100, 1)

    resume_words = set(tokenize(resume_text))
    jd_words = set(tokenize(jd_text))
    return round(len(resume_words & jd_words) / max(len(jd_words), 1) * 100, 1)


def missing_keywords(resume_text: str, jd_text: str) -> list[str]:
    resume_terms = set(tokenize(resume_text))
    jd_terms = top_keywords(jd_text, 30)
    return [term for term in jd_terms if term not in resume_terms][:12]


def call_claude(prompt: str, api_key: str, max_tokens: int = 3000) -> str:
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }

    body = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}],
    }

    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers=headers,
        json=body,
        timeout=90,
    )

    data = response.json()

    if "content" in data:
        return data["content"][0]["text"].strip()

    if "error" in data:
        return f"API Error: {data['error'].get('message', 'Unknown error')}"

    return "Unexpected API response."


def rewrite_resume_plain_text(resume_text: str, jd_text: str, api_key: str) -> str:
    prompt = f"""
You are an expert resume writer and ATS optimization specialist.

Rewrite the candidate resume for the given job description.

Strict rules:
1. Do not invent fake experience, companies, degrees, dates, certifications, or numbers.
2. Keep the candidate's real background.
3. Improve wording professionally.
4. Add missing job-description keywords only where they honestly fit.
5. Use strong action verbs.
6. Keep clear resume headings.
7. Return only the rewritten resume, no explanation.

Original resume:
{resume_text}

Job description:
{jd_text}

Rewrite the resume now:
"""
    return call_claude(prompt, api_key, max_tokens=3000)


def collect_docx_paragraphs(doc):
    items = []
    counter = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            items.append({"id": f"p{counter}", "text": text, "paragraph": para})
            counter += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        items.append({"id": f"p{counter}", "text": text, "paragraph": para})
                        counter += 1

    return items


def replace_paragraph_text_keep_style(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]

    for run in paragraph.runs:
        run.text = ""

    first_run.text = new_text


def rewrite_docx_preserve_format(resume_file, jd_text: str, api_key: str) -> bytes:
    file_bytes = resume_file.getvalue()
    doc = Document(io.BytesIO(file_bytes))

    items = collect_docx_paragraphs(doc)

    paragraph_payload = [
        {"id": item["id"], "text": item["text"]}
        for item in items
    ]

    prompt = f"""
You are an expert resume writer and ATS specialist.

I will give you resume paragraphs with IDs and a job description.

Rewrite each paragraph to tailor the resume to the job description.

Very strict rules:
1. Return valid JSON only.
2. JSON format must be: {{"p1": "rewritten text", "p2": "rewritten text"}}
3. Keep the same paragraph IDs.
4. Do not add new paragraph IDs.
5. Do not delete paragraph IDs.
6. Do not invent fake experience, companies, degrees, dates, or numbers.
7. Keep personal information unchanged.
8. Keep headings short and professional.
9. Add job-description keywords only where they honestly fit.
10. Preserve the meaning of the original paragraph.

Resume paragraphs:
{json.dumps(paragraph_payload, indent=2)}

Job description:
{jd_text}

Return JSON only:
"""

    result = call_claude(prompt, api_key, max_tokens=4000)

    try:
        result = result.strip()
        if result.startswith("```json"):
            result = result.replace("```json", "").replace("```", "").strip()
        elif result.startswith("```"):
            result = result.replace("```", "").strip()

        rewritten_map = json.loads(result)

        for item in items:
            para_id = item["id"]
            if para_id in rewritten_map:
                replace_paragraph_text_keep_style(item["paragraph"], rewritten_map[para_id])

    except Exception:
        fallback_text = rewrite_resume_plain_text(
            "\n".join(item["text"] for item in items),
            jd_text,
            api_key,
        )
        doc = Document()
        doc.add_heading("Tailored Resume", level=1)
        for line in fallback_text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def create_docx_from_text(text: str) -> bytes:
    doc = Document()
    doc.add_heading("Tailored Resume", level=1)

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        lines = block.splitlines()

        if len(lines) == 1 and lines[0].isupper():
            doc.add_heading(lines[0].title(), level=2)
            continue

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.isupper():
                doc.add_heading(line.title(), level=2)
            elif line.startswith("- ") or line.startswith("• "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def chatbot_reply(message: str, resume_text: str, jd_text: str, api_key: str) -> str:
    if not resume_text.strip() or not jd_text.strip():
        return "Please upload or paste both resume and job description first."

    score = calculate_match_score(resume_text, jd_text)
    gaps = missing_keywords(resume_text, jd_text)

    prompt = f"""
You are a professional resume coach.

Resume match score: {score}%
Missing keywords: {", ".join(gaps) or "none"}

Resume:
{resume_text[:1800]}

Job description:
{jd_text[:1200]}

User question:
{message}

Give concise beginner-friendly advice. Do not invent experience.
"""

    return call_claude(prompt, api_key, max_tokens=700)


st.markdown(
    """
    <style>
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }

    .main-header {
        background: #1e293b;
        border: 1px solid #334155;
        border-radius: 10px;
        padding: 24px 28px;
        margin-bottom: 28px;
        box-shadow: 0 4px 18px rgba(0, 0, 0, 0.28);
    }

    .app-title {
        color: #f8fafc !important;
        font-size: 2.1rem;
        font-weight: 800;
        margin: 0 0 8px 0;
        line-height: 1.2;
    }

    .subtle {
        color: #cbd5e1 !important;
        font-size: 1rem;
        margin: 0;
        line-height: 1.5;
    }

    h1, h2, h3 {
        color: #f8fafc !important;
        font-weight: 750 !important;
    }

    textarea {
        font-family: Arial, sans-serif !important;
        font-size: 0.95rem !important;
        border-radius: 8px !important;
    }

    section[data-testid="stFileUploader"] {
        background: #1e293b;
        border: 1px solid #334155;
        border-radius: 10px;
        padding: 16px;
    }

    div[data-testid="stTextArea"] textarea {
        background: #0f172a !important;
        color: #f8fafc !important;
        border: 1px solid #334155 !important;
    }

    .stButton > button {
        background: #2563eb;
        color: white;
        border-radius: 8px;
        border: none;
        font-weight: 700;
        padding: 0.75rem 1rem;
    }

    .stButton > button:hover {
        background: #1d4ed8;
        color: white;
    }

    div[data-testid="stMetric"] {
        background: #1e293b;
        border: 1px solid #334155;
        border-radius: 10px;
        padding: 16px;
    }

    .info-box {
        background: #1e293b;
        border-left: 4px solid #2563eb;
        border-radius: 6px;
        padding: 14px 18px;
        color: #cbd5e1;
        font-size: 0.9rem;
        margin-bottom: 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    """
    <div class="main-header">
        <h1 class="app-title">AI Resume Rewriter for Job Descriptions</h1>
        <p class="subtle">
            Upload your resume and job description. The AI will rewrite the resume professionally,
            improve ATS keyword alignment, and download the result as a Word document.
            If you upload a DOCX resume, the app keeps the original Word format as much as possible.
        </p>
    </div>
    """,
    unsafe_allow_html=True,
)


api_key = st.secrets.get("ANTHROPIC_API_KEY", os.getenv("ANTHROPIC_API_KEY", ""))

if not api_key:
    st.warning(
        "Anthropic API key is missing. For local use, add it in .streamlit/secrets.toml. "
        "For Streamlit Cloud, add it in App Settings > Secrets."
    )


left, right = st.columns([1, 1], gap="large")

with left:
    st.subheader("1. Upload / Paste Resume")
    resume_file = st.file_uploader(
        "Resume file",
        type=["pdf", "docx", "txt"],
        label_visibility="collapsed",
    )
    resume_text = extract_uploaded_text(resume_file)
    resume_text = st.text_area(
        "Resume text",
        value=resume_text,
        height=320,
        placeholder="Upload a resume or paste resume text here...",
    )

with right:
    st.subheader("2. Paste Job Description")
    jd_file = st.file_uploader(
        "Optional JD file",
        type=["txt", "pdf", "docx"],
        label_visibility="collapsed",
    )
    jd_from_file = extract_uploaded_text(jd_file)
    jd_text = st.text_area(
        "Job description",
        value=jd_from_file,
        height=320,
        placeholder="Paste the full job description here...",
    )


st.divider()


if st.button("Rewrite Resume with AI", type="primary", use_container_width=True):
    if not resume_text.strip() or not jd_text.strip():
        st.warning("Please provide both a resume and a job description.")
    elif not api_key.strip():
        st.warning("API key missing. Add ANTHROPIC_API_KEY in Streamlit secrets.")
    else:
        score_before = calculate_match_score(resume_text, jd_text)
        gaps = missing_keywords(resume_text, jd_text)

        with st.spinner("AI is rewriting your resume professionally..."):
            if resume_file is not None and resume_file.name.lower().endswith(".docx"):
                rewritten_docx = rewrite_docx_preserve_format(resume_file, jd_text, api_key)
                rewritten_text = rewrite_resume_plain_text(resume_text, jd_text, api_key)
            else:
                rewritten_text = rewrite_resume_plain_text(resume_text, jd_text, api_key)
                rewritten_docx = create_docx_from_text(rewritten_text)

        if isinstance(rewritten_text, str) and (
            rewritten_text.startswith("API Error") or rewritten_text.startswith("Unexpected")
        ):
            st.error(rewritten_text)
        else:
            score_after = calculate_match_score(rewritten_text, jd_text)

            m1, m2, m3 = st.columns(3)

            with m1:
                st.metric("Match Score Before", f"{score_before}%")

            with m2:
                st.metric(
                    "Match Score After",
                    f"{score_after}%",
                    delta=f"+{round(score_after - score_before, 1)}%",
                )

            with m3:
                st.metric("Missing Keywords Found", f"{len(gaps)}")

            st.progress(min(score_after / 100, 1.0))

            if gaps:
                st.write("Missing / weak keywords from JD:")
                st.write(" ".join(f"`{kw}`" for kw in gaps))

            st.subheader("3. Your Tailored Resume")

            st.markdown(
                """
                <div class="info-box">
                The AI rewrites your resume using the job description while keeping your real experience.
                For DOCX uploads, the downloaded file keeps the original Word format as much as possible.
                Always review before submitting.
                </div>
                """,
                unsafe_allow_html=True,
            )

            st.text_area("Preview of rewritten resume", rewritten_text, height=520)

            st.download_button(
                "Download Tailored Resume as DOCX",
                data=rewritten_docx,
                file_name="tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

else:
    st.info("Upload or paste both inputs, configure your API key, then click Rewrite Resume with AI.")


st.divider()
st.subheader("Resume Coach Chatbot")
st.caption("Ask about match score, missing skills, summary improvement, or resume-vs-JD advice.")

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": "Hi. Upload your resume and job description, then ask me how to improve your resume.",
        }
    ]

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

prompt = st.chat_input("Ask the resume coach...")

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})

    if api_key.strip():
        answer = chatbot_reply(prompt, resume_text, jd_text, api_key)
    else:
        answer = "Please configure ANTHROPIC_API_KEY in Streamlit secrets first."

    st.session_state.messages.append({"role": "assistant", "content": answer})
    st.rerun()